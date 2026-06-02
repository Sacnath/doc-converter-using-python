import os
import threading
import customtkinter as ctk
from tkinter import filedialog, messagebox, PhotoImage
from PIL import Image
import pdfplumber
from docx import Document
import pandas as pd
import pytesseract

# === Setup ===
ctk.set_appearance_mode("Dark")
ctk.set_default_color_theme("blue")

# Set path to Tesseract (you can change this if needed)
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

# === App Window ===
app = ctk.CTk()
app.title("Multi File Converter")
app.geometry("480x600")
app.iconphoto(False, PhotoImage(file="app_icon.png"))
app.resizable(False, False)

frame = ctk.CTkFrame(master=app, corner_radius=12)
frame.pack(padx=30, pady=30, fill="both", expand=True)

ctk.CTkLabel(frame, text="Multi File Converter", font=("Arial", 22, "bold")).pack(pady=(10, 20))

# === Icons ===
icon_size = (24, 24)
icons = {
    "PDF → Word": ctk.CTkImage(Image.open("icons/pdf.png"), size=icon_size),
    "Word → Excel": ctk.CTkImage(Image.open("icons/excel.png"), size=icon_size),
    "Images → PDF": ctk.CTkImage(Image.open("icons/image.png"), size=icon_size),
    "Image → Text": ctk.CTkImage(Image.open("icons/ocr.png"), size=icon_size),
}

# === UI Elements ===
buttons = {}
progress_bar = ctk.CTkProgressBar(master=frame, mode="determinate")
progress_bar.pack(pady=(20, 10), fill="x", padx=20)
progress_bar.set(0)

percent_label = ctk.CTkLabel(frame, text="", font=("Arial", 12))
percent_label.pack()

status_label = ctk.CTkLabel(frame, text="", font=("Arial", 12))
status_label.pack()

# === Conversion Functions ===
def pdf_to_word(pdf_path, output_path):
    doc = Document()
    with pdfplumber.open(pdf_path) as pdf:
        total_pages = len(pdf.pages)
        for i, page in enumerate(pdf.pages):
            text = page.extract_text()
            if text:
                doc.add_paragraph(text)
            update_progress((i + 1) / total_pages)
    doc.save(output_path)
    return True

def word_to_excel(docx_path, output_path):
    from docx import Document as DocxDocument
    doc = DocxDocument(docx_path)
    data = [[para.text.strip()] for para in doc.paragraphs if para.text.strip()]
    pd.DataFrame(data, columns=["Text"]).to_excel(output_path, index=False)
    update_progress(1.0)
    return True

def images_to_pdf(image_paths, output_path):
    imgs = []
    total = len(image_paths)
    for i, path in enumerate(image_paths):
        imgs.append(Image.open(path).convert("RGB"))
        update_progress((i + 1) / total)
    if imgs:
        imgs[0].save(output_path, save_all=True, append_images=imgs[1:])
    return True

def image_to_text(image_path, output_path):
    text = pytesseract.image_to_string(Image.open(image_path))
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(text)
    update_progress(1.0)
    return True

# === Utilities ===
def update_progress(fraction):
    progress_bar.set(fraction)
    percent = int(fraction * 100)
    percent_label.configure(text=f"{percent}%")

def set_loading_state(loading):
    for btn in buttons.values():
        btn.configure(state="disabled" if loading else "normal")
    if loading:
        status_label.configure(text="Converting... Please wait.")
    else:
        status_label.configure(text="")

def show_result(success, mode):
    message = f"{mode} conversion {'successful!' if success else 'failed.'}"
    if success:
        messagebox.showinfo("Success", message)
    else:
        messagebox.showerror("Error", message)

# === Conversion Handler ===
def run_conversion_thread(mode):
    threading.Thread(target=lambda: handle_conversion(mode)).start()

def handle_conversion(mode):
    try:
        set_loading_state(True)

        mapping = {
            "PDF → Word": (["*.pdf"], ".docx", pdf_to_word, False),
            "Word → Excel": (["*.docx"], ".xlsx", word_to_excel, False),
            "Images → PDF": (["*.png","*.jpg","*.jpeg"], "_merged.pdf", images_to_pdf, True),
            "Image → Text": (["*.png","*.jpg","*.jpeg"], ".txt", image_to_text, False)
        }

        extensions, suffix, func, multi = mapping[mode]

        if multi:
            files = filedialog.askopenfilenames(filetypes=[(mode, extensions)])
        else:
            files = filedialog.askopenfilename(filetypes=[(mode, extensions)])
        if not files:
            set_loading_state(False)
            return

        outdir = filedialog.askdirectory()
        if not outdir:
            set_loading_state(False)
            return

        infile = files[0] if isinstance(files, tuple) else files
        base = os.path.splitext(os.path.basename(infile))[0]
        output_file = os.path.join(outdir, base + suffix)

        update_progress(0)  # Reset
        success = func(files if multi else infile, output_file)
        show_result(success, mode)

    except Exception as e:
        messagebox.showerror("Error", str(e))
    finally:
        set_loading_state(False)

# === Buttons ===
btn_info = [
    ("PDF → Word", "PDF → Word"),
    ("Word → Excel", "Word → Excel"),
    ("Images → PDF", "Images → PDF"),
    ("Image → Text", "Image → Text")
]

for label, mode in btn_info:
    btn = ctk.CTkButton(
        master=frame,
        text=label,
        image=icons[mode],
        compound="left",
        command=lambda m=mode: run_conversion_thread(m),
        width=280,
        height=40,
        anchor="w",
        font=("Arial", 14)
    )
    btn.pack(pady=8)
    buttons[mode] = btn

ctk.CTkLabel(frame, text="Created with ❤️ using Python", font=("Arial", 10)).pack(side="bottom", pady=15)

# === Start App ===
app.mainloop()
